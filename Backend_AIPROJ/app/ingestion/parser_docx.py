"""
DOCX parser - extracts text and headings from Word documents.
Requires: python-docx
"""

from pathlib import Path
from typing import Optional, List
from app.core.logging import setup_logger
from .dispatcher import ParsedDocument

logger = setup_logger()


def parse_docx(file_path: str, source: Optional[str] = None) -> ParsedDocument:
    """
    Parse DOCX file using python-docx library.
    
    Args:
        file_path: Path to DOCX file
        source: Optional source identifier
        
    Returns:
        ParsedDocument with extracted text and sections
        
    Raises:
        Exception: If DOCX parsing fails
        ImportError: If python-docx is not installed
    """
    try:
        # Import python-docx (lazy import to avoid dependency issues)
        try:
            from docx import Document
        except ImportError as e:
            raise ImportError(
                "python-docx is not installed. Install with: pip install python-docx"
            ) from e
        
        logger.info(f"Parsing DOCX: {Path(file_path).name}")
        
        path = Path(file_path)
        
        # Load document
        doc = Document(file_path)
        
        # Extract paragraphs and headings
        full_text_parts = []
        sections = []
        current_section = []
        heading_count = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # Check if paragraph is a heading
            if paragraph.style.name.startswith('Heading'):
                heading_count += 1
                
                # Save previous section if exists
                if current_section:
                    sections.append('\n'.join(current_section))
                    current_section = []
                
                # Start new section with heading
                current_section.append(text)
                full_text_parts.append(text)
            else:
                # Regular paragraph
                current_section.append(text)
                full_text_parts.append(text)
        
        # Add last section
        if current_section:
            sections.append('\n'.join(current_section))
        
        # Combine all text
        full_text = '\n\n'.join(full_text_parts)
        
        if not full_text.strip():
            raise ValueError("DOCX file contains no extractable text")
        
        # Build metadata
        metadata = {
            "source": source or path.name,
            "file_name": path.name,
            "paragraph_count": len(doc.paragraphs),
            "heading_count": heading_count,
            "has_sections": len(sections) > 0
        }
        
        # Add core properties if available
        if doc.core_properties:
            props = doc.core_properties
            if props.author:
                metadata["author"] = props.author
            if props.title:
                metadata["title"] = props.title
            if props.created:
                metadata["created"] = props.created.isoformat()
            if props.modified:
                metadata["modified"] = props.modified.isoformat()
        
        logger.info(
            f"DOCX parsed successfully - {len(full_text)} chars, "
            f"{heading_count} headings, {len(sections)} sections"
        )
        
        return ParsedDocument(
            text=full_text,
            sections=sections if len(sections) > 1 else None,
            source_type="document",
            format="docx",
            metadata=metadata
        )
        
    except ImportError:
        raise
    except Exception as e:
        logger.error(f"DOCX parsing failed for {file_path}: {str(e)}")
        raise Exception(f"Failed to parse DOCX: {str(e)}") from e
